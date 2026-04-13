"""
Office & data tools — read and write Excel, Word, PDF, and CSV files.

Dependencies (install as needed):
    pip install openpyxl pypdf python-docx
"""
import csv
import io
from pathlib import Path
from typing import Optional, Union

from .base import BaseTool, ToolResult


# ============================================================
# Excel (XLSX)
# ============================================================

class ReadExcelTool(BaseTool):
    name = "read_excel"
    description = (
        "Read an Excel (.xlsx/.xls) file and return its contents as a list of rows "
        "per sheet. Optionally filter to a specific sheet."
    )

    @property
    def parameters_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Path to the Excel file"},
                "sheet_name": {
                    "type": "string",
                    "description": "Sheet name to read (reads all if omitted)",
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum rows per sheet to return",
                    "default": 500,
                },
            },
            "required": ["path"],
        }

    async def execute(
        self,
        path: str,
        sheet_name: Optional[str] = None,
        max_rows: int = 500,
    ) -> ToolResult:
        try:
            import openpyxl
        except ImportError:
            return ToolResult(success=False, output=None, error="openpyxl not installed. Run: pip install openpyxl")
        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            sheets = [sheet_name] if sheet_name else wb.sheetnames
            result: dict[str, list] = {}
            for name in sheets:
                if name not in wb.sheetnames:
                    return ToolResult(success=False, output=None, error=f"Sheet '{name}' not found")
                ws = wb[name]
                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= max_rows:
                        break
                    rows.append(list(row))
                result[name] = rows
            wb.close()
            return ToolResult(success=True, output=result)
        except Exception as exc:
            return ToolResult(success=False, output=None, error=str(exc))


class WriteExcelTool(BaseTool):
    name = "write_excel"
    description = (
        "Create or overwrite an Excel (.xlsx) file from a dict of "
        "sheet_name → list-of-rows data."
    )

    @property
    def parameters_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Destination .xlsx path"},
                "sheets": {
                    "type": "object",
                    "description": 'Dict of {sheet_name: [[row values], ...]}',
                },
            },
            "required": ["path", "sheets"],
        }

    async def execute(self, path: str, sheets: dict) -> ToolResult:
        try:
            import openpyxl
        except ImportError:
            return ToolResult(success=False, output=None, error="openpyxl not installed.")
        try:
            wb = openpyxl.Workbook()
            wb.remove(wb.active)  # remove default empty sheet
            for sheet_name, rows in sheets.items():
                ws = wb.create_sheet(title=str(sheet_name))
                for row in rows:
                    ws.append([str(c) if c is not None else "" for c in row])
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            wb.save(path)
            return ToolResult(success=True, output=f"Saved Excel file: {path}")
        except Exception as exc:
            return ToolResult(success=False, output=None, error=str(exc))


# ============================================================
# Word (DOCX)
# ============================================================

class ReadWordTool(BaseTool):
    name = "read_word"
    description = "Read a Word (.docx) file and return its text content."

    @property
    def parameters_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Path to the .docx file"},
                "include_tables": {
                    "type": "boolean",
                    "description": "Also extract text from tables",
                    "default": True,
                },
            },
            "required": ["path"],
        }

    async def execute(self, path: str, include_tables: bool = True) -> ToolResult:
        try:
            from docx import Document
        except ImportError:
            return ToolResult(success=False, output=None, error="python-docx not installed. Run: pip install python-docx")
        try:
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            tables_text: list[str] = []
            if include_tables:
                for table in doc.tables:
                    for row in table.rows:
                        tables_text.append("\t".join(cell.text for cell in row.cells))
            content = "\n".join(paragraphs)
            if tables_text:
                content += "\n\n=== Tables ===\n" + "\n".join(tables_text)
            return ToolResult(success=True, output=content)
        except Exception as exc:
            return ToolResult(success=False, output=None, error=str(exc))


class WriteWordTool(BaseTool):
    name = "write_word"
    description = "Create a Word (.docx) file with headings and paragraphs."

    @property
    def parameters_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Destination .docx path"},
                "content": {
                    "type": "array",
                    "description": 'List of {"type": "heading"|"paragraph", "text": "...", "level": 1}',
                    "items": {"type": "object"},
                },
            },
            "required": ["path", "content"],
        }

    async def execute(self, path: str, content: list[dict]) -> ToolResult:
        try:
            from docx import Document
        except ImportError:
            return ToolResult(success=False, output=None, error="python-docx not installed.")
        try:
            doc = Document()
            for block in content:
                btype = block.get("type", "paragraph")
                text = block.get("text", "")
                level = int(block.get("level", 1))
                if btype == "heading":
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(text)
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(path)
            return ToolResult(success=True, output=f"Saved Word file: {path}")
        except Exception as exc:
            return ToolResult(success=False, output=None, error=str(exc))


# ============================================================
# PDF
# ============================================================

class ReadPDFTool(BaseTool):
    name = "read_pdf"
    description = "Extract text from a PDF file, page by page."

    @property
    def parameters_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Path to the PDF"},
                "pages": {
                    "type": "string",
                    "description": 'Page range, e.g. "1-5" or "all"',
                    "default": "all",
                },
            },
            "required": ["path"],
        }

    async def execute(self, path: str, pages: str = "all") -> ToolResult:
        try:
            from pypdf import PdfReader
        except ImportError:
            return ToolResult(success=False, output=None, error="pypdf not installed. Run: pip install pypdf")
        try:
            reader = PdfReader(path)
            total = len(reader.pages)

            if pages == "all":
                indices = range(total)
            else:
                parts = pages.split("-")
                start = int(parts[0]) - 1
                end = int(parts[1]) if len(parts) > 1 else start + 1
                indices = range(max(0, start), min(total, end))

            text_parts = []
            for i in indices:
                page_text = reader.pages[i].extract_text() or ""
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

            return ToolResult(success=True, output="\n\n".join(text_parts))
        except Exception as exc:
            return ToolResult(success=False, output=None, error=str(exc))


# ============================================================
# CSV
# ============================================================

class ReadCSVTool(BaseTool):
    name = "read_csv"
    description = "Read a CSV file and return rows as a list of dicts (first row = headers)."

    @property
    def parameters_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "delimiter": {"type": "string", "default": ","},
                "max_rows": {"type": "integer", "default": 1000},
            },
            "required": ["path"],
        }

    async def execute(
        self, path: str, delimiter: str = ",", max_rows: int = 1000
    ) -> ToolResult:
        try:
            rows: list[dict] = []
            with open(path, newline="", encoding="utf-8", errors="replace") as fh:
                reader = csv.DictReader(fh, delimiter=delimiter)
                for i, row in enumerate(reader):
                    if i >= max_rows:
                        break
                    rows.append(dict(row))
            return ToolResult(success=True, output=rows)
        except Exception as exc:
            return ToolResult(success=False, output=None, error=str(exc))


class WriteCSVTool(BaseTool):
    name = "write_csv"
    description = "Write a list of dicts (or list of lists) to a CSV file."

    @property
    def parameters_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "rows": {
                    "type": "array",
                    "description": "List of row dicts or list of lists",
                    "items": {},
                },
                "headers": {
                    "type": "array",
                    "description": "Column headers (auto-detected from dicts if omitted)",
                    "items": {"type": "string"},
                },
                "delimiter": {"type": "string", "default": ","},
            },
            "required": ["path", "rows"],
        }

    async def execute(
        self,
        path: str,
        rows: list,
        headers: Optional[list[str]] = None,
        delimiter: str = ",",
    ) -> ToolResult:
        try:
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            with open(path, "w", newline="", encoding="utf-8") as fh:
                if rows and isinstance(rows[0], dict):
                    fieldnames = headers or list(rows[0].keys())
                    writer = csv.DictWriter(fh, fieldnames=fieldnames, delimiter=delimiter)
                    writer.writeheader()
                    writer.writerows(rows)
                else:
                    writer_plain = csv.writer(fh, delimiter=delimiter)
                    if headers:
                        writer_plain.writerow(headers)
                    writer_plain.writerows(rows)
            return ToolResult(success=True, output=f"Saved CSV: {path} ({len(rows)} rows)")
        except Exception as exc:
            return ToolResult(success=False, output=None, error=str(exc))
